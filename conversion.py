"""Conversion pipeline. Wraps Stages 1, 3, and 4 of the spec.

Stage 2 (cleanups) is in `cleanups.py` and called from here.
"""
from __future__ import annotations

import json
import shutil
import time
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import Optional

import pypandoc

import cleanups
import db
from config import CONTENT_DIR, VERSIONS_KEEP


# ---------- helpers ----------

def article_dir(journal_slug: str, issue_slug: Optional[str], article_slug: str) -> Path:
    if issue_slug:
        d = CONTENT_DIR / "journals" / journal_slug / "issues" / issue_slug / "articles" / article_slug
    else:
        d = CONTENT_DIR / "journals" / journal_slug / "_unfiled" / article_slug
    d.mkdir(parents=True, exist_ok=True)
    (d / "assets").mkdir(exist_ok=True)
    (d / ".versions").mkdir(exist_ok=True)
    return d


def issue_dir(journal_slug: str, issue_slug: str) -> Path:
    d = CONTENT_DIR / "journals" / journal_slug / "issues" / issue_slug
    d.mkdir(parents=True, exist_ok=True)
    (d / "articles").mkdir(exist_ok=True)
    return d


def issue_slug_for(volume, issue_number, year) -> str:
    return f"v{volume}-n{issue_number}-{year}"


def move_article_to_issue(article_path: Path, journal_slug: str, issue_slug: str, article_slug: str) -> Path:
    """Move an article directory to its issue location. Returns new path."""
    import shutil
    dest = CONTENT_DIR / "journals" / journal_slug / "issues" / issue_slug / "articles" / article_slug
    if dest.exists():
        raise FileExistsError(f"Destination already exists: {dest}")
    dest.parent.mkdir(parents=True, exist_ok=True)
    shutil.move(str(article_path), str(dest))
    return dest


def move_article_to_unfiled(article_path: Path, journal_slug: str, article_slug: str) -> Path:
    """Move an article directory back to _unfiled. Returns new path."""
    import shutil
    dest = CONTENT_DIR / "journals" / journal_slug / "_unfiled" / article_slug
    if dest.exists():
        raise FileExistsError(f"Destination already exists: {dest}")
    dest.parent.mkdir(parents=True, exist_ok=True)
    shutil.move(str(article_path), str(dest))
    return dest


def write_issue_yaml(journal_slug: str, issue_slug: str, payload: dict):
    """Persist issue-level metadata to `_issue.yaml` alongside the issue directory."""
    import yaml
    d = issue_dir(journal_slug, issue_slug)
    (d / "_issue.yaml").write_text(
        yaml.safe_dump(payload, sort_keys=False, allow_unicode=True, width=10_000),
        encoding="utf-8",
    )


def template_dir(journal_slug: str) -> Path:
    return CONTENT_DIR / "journals" / journal_slug / "template"


def _append_log(article_path: Path, header: str, body: str):
    log_file = article_path / "conversion.log"
    stamp = datetime.now().isoformat(timespec="seconds")
    with log_file.open("a", encoding="utf-8") as f:
        f.write(f"\n=== {stamp}  {header} ===\n{body}\n")


def _pandoc_version() -> str:
    try:
        return pypandoc.get_pandoc_version()
    except Exception:
        return "unknown"


def _snapshot_version(article_path: Path):
    """Save current article.md to .versions/ before overwrite. Keep last N."""
    src = article_path / "article.md"
    if not src.exists():
        return
    versions_dir = article_path / ".versions"
    versions_dir.mkdir(exist_ok=True)
    stamp = datetime.now().strftime("%Y%m%dT%H%M%S")
    shutil.copy2(src, versions_dir / f"article-{stamp}.md")
    snaps = sorted(versions_dir.glob("article-*.md"))
    while len(snaps) > VERSIONS_KEEP:
        snaps.pop(0).unlink(missing_ok=True)


# ---------- Stage 1: DOCX ingest ----------

@dataclass
class IngestResult:
    raw_md_path: Path
    log: str
    has_tracked_changes: bool


def ingest_docx(
    docx_path: Path,
    article_path: Path,
    accept_track_changes: bool = True,
) -> IngestResult:
    """Run Pandoc on the DOCX, extract media, write `article-raw.md`.

    Before invoking Pandoc, scan the docx for a `Title`-styled paragraph
    (Pandoc does not promote that style to a markdown heading). When
    found, inject it as Pandoc metadata so the raw markdown opens with
    a YAML front matter block.
    """
    article_path.mkdir(parents=True, exist_ok=True)
    assets = article_path / "assets"
    assets.mkdir(exist_ok=True)

    source_copy = article_path / "source.docx"
    if Path(docx_path).resolve() != source_copy.resolve():
        shutil.copy2(docx_path, source_copy)

    raw_md = article_path / "article-raw.md"

    docx_title = _extract_docx_title(source_copy)

    extra = [
        "--wrap=none",
        "--markdown-headings=atx",
        f"--extract-media={assets}",
        f"--track-changes={'accept' if accept_track_changes else 'reject'}",
    ]

    t0 = time.time()
    pypandoc.convert_file(
        str(source_copy),
        to="markdown",
        format="docx",
        outputfile=str(raw_md),
        extra_args=extra,
    )
    dt = time.time() - t0

    has_tracked = _docx_has_tracked_changes(source_copy)

    body = (
        f"pandoc version: {_pandoc_version()}\n"
        f"input: {source_copy.name}\n"
        f"output: {raw_md.name}\n"
        f"extract-media: {assets.relative_to(article_path)}\n"
        f"track-changes: {'accept' if accept_track_changes else 'reject'}\n"
        f"tracked-changes-present: {has_tracked}\n"
        f"docx-title-detected: {docx_title or '(none)'}\n"
        f"elapsed: {dt:.2f}s\n"
    )
    _append_log(article_path, "Stage 1: DOCX ingest", body)

    return IngestResult(raw_md_path=raw_md, log=body, has_tracked_changes=has_tracked)


def _extract_docx_title(docx_path: Path) -> Optional[str]:
    """Pull the title from the docx Title-styled paragraph or core properties."""
    try:
        import docx as _docx
        d = _docx.Document(str(docx_path))
    except Exception:
        return None

    core_title = (d.core_properties.title or "").strip()
    if core_title:
        return core_title

    for p in d.paragraphs[:20]:
        style = (p.style.name if p.style else "") or ""
        text = p.text.strip()
        if not text:
            continue
        if style.lower() in {"title", "document title", "subtitle"} and style.lower() != "subtitle":
            return text.replace("\n", " ").strip()
        if style.lower() == "title":
            return text
    return None


def _docx_has_tracked_changes(docx_path: Path) -> bool:
    """Quick heuristic: unzip docx and check for w:ins / w:del tags."""
    import zipfile
    try:
        with zipfile.ZipFile(docx_path) as z:
            if "word/document.xml" not in z.namelist():
                return False
            xml = z.read("word/document.xml").decode("utf-8", errors="replace")
            return ("<w:ins " in xml) or ("<w:del " in xml)
    except Exception:
        return False


# ---------- Stage 2 orchestration ----------

def run_cleanups(article_path: Path, issue_metadata: Optional[dict] = None) -> Path:
    """Read article-raw.md, apply cleanups, write article.md. Snapshot first.

    Pulls the docx title from source.docx if present and merges it into
    the YAML front matter (Pandoc's markdown writer does not promote
    Word's Title style to anything recoverable from the body alone).
    """
    raw = (article_path / "article-raw.md").read_text(encoding="utf-8")

    extra_metadata = dict(issue_metadata or {})
    source_docx = article_path / "source.docx"
    if source_docx.exists() and "title" not in extra_metadata:
        title = _extract_docx_title(source_docx)
        if title:
            extra_metadata["title"] = title

    cleaned, log = cleanups.run_all(raw, issue_metadata=extra_metadata)

    _snapshot_version(article_path)
    (article_path / "article.md").write_text(cleaned, encoding="utf-8")
    _append_log(article_path, "Stage 2: cleanups", log.render())
    return article_path / "article.md"


def save_markdown(article_path: Path, new_text: str, note: str = "manual edit"):
    """Editor save. Snapshots the previous version."""
    _snapshot_version(article_path)
    (article_path / "article.md").write_text(new_text, encoding="utf-8")
    _append_log(article_path, "Stage 3: editor save", f"note: {note}\nbytes: {len(new_text)}")


# ---------- metadata helpers ----------

# Canonical YAML field order. Keys outside this list are appended after.
_FIELD_ORDER = (
    "title", "subtitle",
    "author",
    "abstract",
    "keywords",
    "short-title", "short-authors", "footer",
    "doi",
    "journal", "issn", "volume", "issue", "year",
    "start-page", "end-page",
    "submitted-date", "accepted-date", "published-date",
    "copyright",
    "status",
)


def read_article_metadata(article_path: Path) -> tuple[dict, str]:
    """Return (front_matter_dict, body_string) from article.md.

    If no YAML front matter is present, returns ({}, full_text).
    """
    import yaml
    md_path = article_path / "article.md"
    text = md_path.read_text(encoding="utf-8") if md_path.exists() else ""
    if not text.startswith("---\n"):
        return {}, text
    end = text.find("\n---", 4)
    if end == -1:
        return {}, text
    raw = text[4:end]
    body_start = end + len("\n---")
    if text[body_start:body_start + 1] == "\n":
        body_start += 1
    try:
        fm = yaml.safe_load(raw) or {}
    except Exception:
        fm = {}
    return fm, text[body_start:]


# Fields that must remain single-line strings (no embedded newlines).
_SINGLE_LINE_FIELDS = frozenset({
    "title", "subtitle", "short-title", "short-authors", "footer",
    "doi", "journal", "issn", "status", "copyright",
})


def _sanitize_scalar(s):
    """Collapse interior whitespace in a string scalar so it serializes as a
    plain single-line YAML value. Browsers and copy-paste can occasionally
    introduce stray newlines or tabs into form values; if any reach
    safe_dump in fields meant to be single-line, the resulting YAML can
    parse ambiguously. Defensive normalization at the boundary."""
    if not isinstance(s, str):
        return s
    return " ".join(s.split())


def write_article_metadata(article_path: Path, fm: dict, body: str | None = None):
    """Write article.md with the given front matter and body. Snapshots first.

    If `body` is None, preserves the current body. Field order is canonical
    (title first, then authors, etc.) so files diff cleanly between saves.
    Single-line scalar fields are sanitized (interior whitespace collapsed)
    to avoid round-trip corruption.
    """
    import yaml
    if body is None:
        _, body = read_article_metadata(article_path)

    ordered: dict = {}
    for k in _FIELD_ORDER:
        if k in fm and fm[k] not in (None, "", []):
            v = fm[k]
            if k in _SINGLE_LINE_FIELDS:
                v = _sanitize_scalar(v)
            ordered[k] = v
    for k, v in fm.items():
        if k not in ordered and v not in (None, "", []):
            if k in _SINGLE_LINE_FIELDS:
                v = _sanitize_scalar(v)
            ordered[k] = v

    # Also sanitize author name/affiliation/orcid as single-line.
    if "author" in ordered and isinstance(ordered["author"], list):
        for a in ordered["author"]:
            if isinstance(a, dict):
                for k in ("name", "affiliation", "orcid", "email"):
                    if k in a and isinstance(a[k], str):
                        a[k] = _sanitize_scalar(a[k])

    yaml_text = yaml.safe_dump(
        ordered, sort_keys=False, allow_unicode=True, width=10_000, default_flow_style=False
    )

    # Defensive round-trip check: if we cannot re-parse what we just wrote,
    # something is wrong with the input dict, not the file. Raise so the
    # caller can surface the issue rather than persist corrupt YAML.
    try:
        yaml.safe_load(yaml_text)
    except yaml.YAMLError as exc:
        raise ValueError(f"Refusing to write unparseable YAML: {exc}") from exc

    out_text = f"---\n{yaml_text}---\n\n{body.lstrip(chr(10))}"

    _snapshot_version(article_path)
    (article_path / "article.md").write_text(out_text, encoding="utf-8")
    _append_log(
        article_path,
        "Stage 3: metadata save",
        f"fields: {', '.join(ordered.keys())}\nbytes: {len(out_text)}",
    )


# ---------- Stage 4: render ----------

@dataclass
class RenderResult:
    html_path: Optional[Path]
    pdf_path: Optional[Path]
    errors: list


def render_html(article_path: Path, journal_slug: str) -> Path:
    md = article_path / "article.md"
    out = article_path / "article.html"
    tpl = template_dir(journal_slug)
    template = tpl / "article.html.j2"
    css = tpl / "article.css"
    lua_filter = tpl / "lics-filter.lua"

    extra = [
        "--standalone",
        "--section-divs",
        f"--template={template}",
        f"--css={css.name}",
    ]
    if lua_filter.exists():
        extra.append(f"--lua-filter={lua_filter}")

    if css.exists():
        shutil.copy2(css, article_path / css.name)

    pypandoc.convert_file(
        str(md),
        to="html5",
        format="markdown+yaml_metadata_block",
        outputfile=str(out),
        extra_args=extra,
    )
    return out


def render_pdf(article_path: Path, journal_slug: str) -> Path:
    """Render PDF via Pandoc (Typst template) + typst-py compile.

    Pandoc reads article.md, substitutes YAML front matter into the
    journal's Pandoc Typst template, and writes article.typ. The typst
    Python package then compiles that to article.pdf.
    """
    md = article_path / "article.md"
    out = article_path / "article.pdf"
    tpl = template_dir(journal_slug)
    typ_template = tpl / "article.typ"
    lua_filter = tpl / "lics-filter.lua"
    typst_input = article_path / "article.typ"

    extra = [f"--template={typ_template}"]
    if lua_filter.exists():
        extra.append(f"--lua-filter={lua_filter}")
    pypandoc.convert_file(
        str(md),
        to="typst",
        format="markdown+yaml_metadata_block",
        outputfile=str(typst_input),
        extra_args=extra,
    )

    import typst as typst_lib
    typst_lib.compile(str(typst_input), output=str(out), root=str(CONTENT_DIR))
    return out


def render_all(article_path: Path, journal_slug: str) -> RenderResult:
    errors = []
    html_path = pdf_path = None
    try:
        html_path = render_html(article_path, journal_slug)
    except Exception as exc:
        errors.append(f"HTML: {type(exc).__name__}: {exc}")
    try:
        pdf_path = render_pdf(article_path, journal_slug)
    except Exception as exc:
        errors.append(f"PDF: {type(exc).__name__}: {exc}")

    body = json.dumps(
        {
            "html": str(html_path) if html_path else None,
            "pdf": str(pdf_path) if pdf_path else None,
            "errors": errors,
        },
        indent=2,
    )
    _append_log(article_path, "Stage 4: render", body)
    return RenderResult(html_path=html_path, pdf_path=pdf_path, errors=errors)


# ---------- DB integration ----------

def record_conversion(article_id: int, source_format: str, notes: str, success: bool = True):
    db.execute(
        "INSERT INTO conversions (article_id, source_format, pandoc_version, notes, success) "
        "VALUES (?, ?, ?, ?, ?)",
        (article_id, source_format, _pandoc_version(), notes, 1 if success else 0),
    )


# ---------- issue assembly ----------

def _pdf_page_count(pdf_path: Path) -> int:
    import pypdfium2 as pdfium
    doc = pdfium.PdfDocument(str(pdf_path))
    try:
        return len(doc)
    finally:
        doc.close()


@dataclass
class AssemblyResult:
    article_pages: list  # list of (article_id, start_page, end_page)
    issue_pdf_path: Path
    total_pages: int
    errors: list


def assemble_issue(issue_id: int) -> AssemblyResult:
    """Build the combined-issue PDF with continuous pagination.

    For each article in order:
      1. Write start-page into article.md YAML
      2. Re-render the PDF (Typst start-page-val shifts the page counter)
      3. Count the rendered pages, store start/end in DB

    Then render the issue cover (issue.typ) and concatenate cover + articles
    into issues/<slug>/issue.pdf.
    """
    import json
    from pypdf import PdfWriter, PdfReader

    issue_row = db.query_one(
        "SELECT i.*, j.slug AS journal_slug, j.name AS journal_name, j.issn AS journal_issn "
        "FROM issues i JOIN journals j ON i.journal_id = j.id WHERE i.id = ?",
        (issue_id,),
    )
    if not issue_row:
        raise ValueError(f"Issue {issue_id} not found")
    issue = dict(issue_row)

    journal_row = db.query_one("SELECT * FROM journals WHERE id = ?", (issue["journal_id"],))
    journal = dict(journal_row) if journal_row else {}

    article_rows = db.query_all(
        "SELECT * FROM articles WHERE issue_id = ? "
        "ORDER BY COALESCE(order_in_issue, 999999), updated_at",
        (issue_id,),
    )
    if not article_rows:
        raise ValueError("Issue has no articles to assemble")
    articles_all = [dict(r) for r in article_rows]
    articles = [a for a in articles_all if (a.get("kind") or "article") != "editorial"]
    editor_intro = next((a for a in articles_all if (a.get("kind") or "article") == "editorial"), None)

    errors: list = []
    article_pages: list = []
    cumulative = 0
    article_pdf_paths: list = []

    for art in articles:
        apath = Path(art["project_path"])
        start_page = cumulative + 1

        fm, body = read_article_metadata(apath)
        fm["start-page"] = start_page
        write_article_metadata(apath, fm, body)

        try:
            pdf_path = render_pdf(apath, issue["journal_slug"])
        except Exception as exc:
            errors.append(f"render {art['slug']}: {type(exc).__name__}: {exc}")
            continue

        page_count = _pdf_page_count(pdf_path)
        end_page = start_page + page_count - 1
        db.execute(
            "UPDATE articles SET start_page = ?, end_page = ?, updated_at = CURRENT_TIMESTAMP "
            "WHERE id = ?",
            (start_page, end_page, art["id"]),
        )

        article_pages.append((art["id"], start_page, end_page))
        article_pdf_paths.append(pdf_path)
        cumulative = end_page

    # Refresh article rows so start_page is included for the ToC.
    refreshed_rows = db.query_all(
        "SELECT * FROM articles WHERE issue_id = ? "
        "ORDER BY COALESCE(order_in_issue, 999999), updated_at",
        (issue_id,),
    )
    refreshed = [dict(r) for r in refreshed_rows]
    articles_for_toc = [a for a in refreshed if (a.get("kind") or "article") != "editorial"]

    toc_entries = [{
        "title": a["title"],
        "authors": "",  # filled in by front-matter renderer from each article's YAML
        "start_page": a.get("start_page"),
        "end_page": a.get("end_page"),
        "section": a.get("section") or "ARTICLES",
    } for a in articles_for_toc]

    issue_dir_path = issue_dir(issue["journal_slug"], issue_slug_for(issue["volume"], issue["issue_number"], issue["year"]))
    (issue_dir_path / "issue-toc.json").write_text(
        json.dumps(toc_entries, indent=2), encoding="utf-8"
    )

    front_pdf = render_front_matter(issue, journal, articles_for_toc, editor_intro, issue_dir_path)

    issue_pdf = issue_dir_path / "issue.pdf"
    writer = PdfWriter()
    for pdf_path in [front_pdf] + article_pdf_paths:
        reader = PdfReader(str(pdf_path))
        for page in reader.pages:
            writer.add_page(page)
    with issue_pdf.open("wb") as f:
        writer.write(f)

    return AssemblyResult(
        article_pages=article_pages,
        issue_pdf_path=issue_pdf,
        total_pages=cumulative,
        errors=errors,
    )


def _authors_inline(authors) -> str:
    if not authors:
        return ""
    names: list = []
    for a in authors:
        if isinstance(a, dict):
            names.append(a.get("name", ""))
        else:
            names.append(str(a))
    names = [n for n in names if n]
    if len(names) <= 2:
        return ", ".join(names)
    return ", ".join(names[:-1]) + ", and " + names[-1]


def _md_to_typst_fragment(md_text: Optional[str]) -> str:
    """Convert a Markdown string to a Typst content fragment via Pandoc.
    Returns an empty string if the input is empty."""
    if not md_text or not md_text.strip():
        return ""
    return pypandoc.convert_text(
        md_text, to="typst", format="markdown",
        extra_args=["--wrap=preserve"],
    )


def _read_editor_intro_body(article: Optional[dict]) -> str:
    """Pull the body markdown from an editor's-intro article, stripping
    YAML front matter."""
    if not article:
        return ""
    md_path = Path(article["project_path"]) / "article.md"
    if not md_path.exists():
        return ""
    text = md_path.read_text(encoding="utf-8")
    if text.startswith("---"):
        end = text.find("\n---", 4)
        if end != -1:
            return text[end + len("\n---"):].lstrip()
    return text


def _toc_section_order(journal: dict) -> list[str]:
    """Return the journal's configured ToC section labels in order, or
    a sensible default."""
    raw = journal.get("toc_sections_json")
    if raw:
        try:
            import json
            sections = json.loads(raw)
            if isinstance(sections, list) and sections:
                return [str(s) for s in sections]
        except Exception:
            pass
    return ["ARTICLES"]


def render_front_matter(issue: dict, journal: dict, articles: list, editor_intro: Optional[dict], out_dir: Path) -> Path:
    """Render the issue's full front matter (cover, editorial team,
    board + credits, mission statement, editor's introduction, ToC) as
    a single roman-paginated PDF.

    Articles with kind=editorial are excluded from the ToC. The editor's
    intro itself comes in via the `editor_intro` argument (the article
    record). Returns the path to front_matter.pdf.
    """
    import typst as typst_lib

    vol = issue["volume"]
    num = issue["issue_number"]
    year = issue["year"]
    season = issue.get("header_season") or ""
    title = issue.get("title") or ""
    journal_name = journal.get("name") or issue.get("journal_name") or ""
    short_name = journal.get("short_name") or _short_journal_name(journal_name)
    issn = issue.get("journal_issn") or journal.get("issn") or ""

    header_template = journal.get("header_label_template") or "*{short_name}* {volume}.{issue} / {season}"
    header_label = header_template.format(
        short_name=short_name, name=journal_name,
        volume=vol, issue=num, year=year, season=season,
    ).strip()
    if header_label.endswith("/"):
        header_label = header_label[:-1].strip()

    wordmark_block = _typst_wordmark_block(journal.get("wordmark_image_path"), out_dir, short_name)

    team_typst = _md_to_typst_fragment(journal.get("editorial_team_md"))
    board_typst = _md_to_typst_fragment(journal.get("editorial_board_md"))
    credit_typst = _md_to_typst_fragment(journal.get("financial_credit_md"))
    mission_typst = _md_to_typst_fragment(journal.get("mission_statement_md"))
    intro_typst = _md_to_typst_fragment(_read_editor_intro_body(editor_intro))

    # Build ToC entries grouped by section, in journal-configured order
    section_order = _toc_section_order(journal)
    toc_by_section: dict[str, list] = {label: [] for label in section_order}
    for art in articles:
        if (art.get("kind") or "article") == "editorial":
            continue
        section = (art.get("section") or section_order[0]).strip()
        toc_by_section.setdefault(section, []).append(art)

    # Compose ToC Typst
    toc_blocks = []
    for label in list(toc_by_section.keys()):
        entries = toc_by_section.get(label) or []
        if not entries:
            continue
        toc_blocks.append(
            f'#v(0.5em)\n'
            f'#grid(columns: (1fr, auto), gutter: 0.5em,\n'
            f'  text(size: 11pt, tracking: 0.05em, {_typst_str(label.upper())}),\n'
            f'  line(length: 100%, stroke: 0.5pt + rgb("#1a1612")),\n'
            f')\n'
            f'#v(0.5em)\n'
        )
        for art in entries:
            md_path = Path(art["project_path"]) / "article.md"
            fm = {}
            try:
                if md_path.exists():
                    fm, _ = read_article_metadata(Path(art["project_path"]))
            except Exception:
                fm = {}
            title_text = fm.get("title") or art["title"]
            authors_text = _authors_inline_with_affil(fm.get("author") or [])
            start = art.get("start_page") or "?"
            toc_blocks.append(
                f'#grid(columns: (3em, 1fr), gutter: 0.4em,\n'
                f'  align(left, text(size: 10pt, {_typst_str(str(start))})),\n'
                f'  block(width: 100%, [\n'
                f'    #text(size: 10pt, {_typst_str(title_text)})\n'
                f'    {("#linebreak()" + chr(10) + "    #text(size: 10pt, style: " + chr(34) + "italic" + chr(34) + ", " + _typst_str(authors_text) + ")") if authors_text else ""}\n'
                f'  ]),\n'
                f')\n'
                f'#v(0.6em)\n'
            )
    toc_typst = "".join(toc_blocks)

    intro_heading = f"EDITORS' INTRODUCTION TO ISSUE {vol}.{num}"

    typst_doc = f"""
#let ink = rgb("#1a1612")
#let ink-soft = rgb("#4a4137")
#let rule-color = rgb("#b6a98c")
#let header-label = [{_typst_inline_md(header_label)}]

#set page(
  paper: "us-letter",
  width: 6in,
  height: 9in,
  margin: (top: 0.85in, bottom: 0.95in, left: 0.75in, right: 0.75in),
  numbering: "I",
  number-align: right,
  header: context {{
    let p = counter(page).at(here()).first()
    if p == 1 {{ return [] }}
    align(center, text(size: 9pt, fill: ink, style: "italic", header-label))
  }},
)

#set text(font: ("EB Garamond", "Garamond", "Georgia"), size: 11pt, fill: ink, hyphenate: false)
#set par(justify: true, leading: 0.65em, first-line-indent: 0pt)

// Tables: thin rules, generous padding, EB Garamond throughout
#show table.cell: set text(size: 10.5pt)
#set table(stroke: 0.5pt + rule-color, inset: 8pt)

// ============ PAGE I: COVER ============

#align(center, text(size: 9pt, style: "italic", fill: ink-soft, header-label))
{wordmark_block}

#v(1fr)

#align(right, block(width: auto, {{
  set par(first-line-indent: 0pt)
  text(size: 22pt, weight: 600, "{year}")
  linebreak()
  text(size: 14pt, "Volume {vol}, Number {num}")
  {f'linebreak(); v(0.4em); text(size: 11pt, style: "italic", fill: ink-soft, {_typst_str(title)})' if title else ''}
}}))

#v(0.4in)

// ============ PAGE II: EDITORIAL TEAM ============
{("#pagebreak()" + chr(10) + chr(10) + "#v(0.5in)" + chr(10) + "#align(center, text(size: 14pt, tracking: 0.18em, " + _typst_str(journal_name.upper()) + "))" + chr(10) + "#v(1em)" + chr(10) + chr(10) + team_typst) if team_typst else ""}

// ============ PAGE III: EDITORIAL BOARD ============
{("#pagebreak()" + chr(10) + chr(10) + "#v(0.4in)" + chr(10) + "#align(center, text(size: 14pt, tracking: 0.18em, " + chr(34) + "EDITORIAL BOARD" + chr(34) + "))" + chr(10) + "#v(1em)" + chr(10) + chr(10) + board_typst) if board_typst else ""}

{("#v(1.5em)" + chr(10) + "#align(center, block(width: 70%, [" + chr(10) + "  #set par(first-line-indent: 0pt, justify: false, leading: 0.6em)" + chr(10) + "  #set text(size: 10pt, style: " + chr(34) + "italic" + chr(34) + ", fill: ink-soft)" + chr(10) + "  " + credit_typst + chr(10) + "]))") if credit_typst else ""}

// ============ MISSION STATEMENT ============
{("#pagebreak()" + chr(10) + chr(10) + "#v(0.4in)" + chr(10) + "#align(center, text(size: 14pt, tracking: 0.18em, [#text(style: " + chr(34) + "italic" + chr(34) + ", " + chr(34) + short_name + chr(34) + ") MISSION STATEMENT]))" + chr(10) + "#v(1em)" + chr(10) + chr(10) + "#set par(first-line-indent: 1.4em, justify: true)" + chr(10) + mission_typst) if mission_typst else ""}

// ============ EDITORS' INTRODUCTION ============
{("#pagebreak()" + chr(10) + chr(10) + "#v(0.4in)" + chr(10) + "#align(center, text(size: 14pt, tracking: 0.18em, " + _typst_str(intro_heading) + "))" + chr(10) + "#v(1em)" + chr(10) + chr(10) + "#set par(first-line-indent: 1.4em, justify: true)" + chr(10) + intro_typst) if intro_typst else ""}

// ============ TABLE OF CONTENTS ============
#pagebreak()

#v(0.4in)
#align(center, text(size: 14pt, tracking: 0.18em, "CONTENTS"))
#v(1.5em)

#set par(first-line-indent: 0pt, justify: false)
{toc_typst}
"""

    out_path = out_dir / "_front_matter.typ"
    out_pdf = out_dir / "_front_matter.pdf"
    out_path.write_text(typst_doc, encoding="utf-8")
    typst_lib.compile(str(out_path), output=str(out_pdf), root=str(CONTENT_DIR))
    return out_pdf


def _authors_inline_with_affil(authors) -> str:
    if not authors:
        return ""
    parts = []
    for a in authors:
        if isinstance(a, dict):
            name = a.get("name", "")
            aff = a.get("affiliation", "")
            if aff:
                parts.append(f"{name} — {aff}")
            else:
                parts.append(name)
        else:
            parts.append(str(a))
    return "; ".join(p for p in parts if p)


def _render_issue_cover(issue, toc_entries: list, out_dir: Path) -> Path:
    """Render the issue's title page (and a basic ToC for now) as a
    front-matter PDF with roman pagination.

    Layout: top-centered italic running label, large wordmark (image if
    journal.wordmark_image_path is set, text fallback otherwise),
    journal subtitle, year/volume/number block bottom-right, roman page
    counter. Subsequent front-matter sections (editorial team, board,
    mission, intro, ToC) will be added in follow-up commits.
    """
    import typst as typst_lib

    journal_name = issue["journal_name"]
    title = issue["title"] or ""
    vol = issue["volume"]
    num = issue["issue_number"]
    year = issue["year"]
    season = (issue.get("header_season") if isinstance(issue, dict) else issue["header_season"]) or ""

    journal_row = db.query_one(
        "SELECT wordmark_image_path, header_label_template, name, short_name FROM journals WHERE id = ?",
        (issue["journal_id"],),
    )
    journal = dict(journal_row) if journal_row else {}

    short_name = journal.get("short_name") or _short_journal_name(journal.get("name") or journal_name)
    header_template = journal.get("header_label_template") or "*{short_name}* {volume}.{issue} / {season}"
    header_label = header_template.format(
        short_name=short_name,
        name=journal.get("name") or journal_name,
        volume=vol, issue=num, year=year,
        season=season,
    ).strip()
    if header_label.endswith("/"):
        header_label = header_label[:-1].strip()

    wordmark_block = _typst_wordmark_block(journal.get("wordmark_image_path"), out_dir, short_name)

    issue_label_lines = [str(year), f"Volume {vol}, Number {num}"]
    if title:
        issue_label_lines.insert(0, title)

    toc_typst = ""
    for e in toc_entries:
        t = _typst_str(e["title"])
        a = _typst_str(e["authors"])
        start = e["start_page"]
        toc_typst += (
            f'grid(columns: (1fr, auto), '
            f'[#text(weight: 500, {t}) \\\n  #text(style: "italic", fill: rgb("#4a4137"), {a})], '
            f'align(right, text({_typst_str(str(start))})))\n'
            f'v(0.5em)\n'
        )

    cover_typst = f"""
#set page(
  paper: "us-letter",
  width: 6in,
  height: 9in,
  margin: (top: 0.85in, bottom: 0.95in, left: 0.75in, right: 0.75in),
  numbering: "I",
  number-align: right,
  header: align(center, text(
    size: 9pt,
    fill: rgb("#1a1612"),
    style: "italic",
    [{_typst_inline_md(header_label)}]
  )),
)

#set text(font: ("EB Garamond", "Garamond", "Georgia"), size: 11pt, fill: rgb("#1a1612"), hyphenate: false)
#set par(justify: false, leading: 0.65em)

{wordmark_block}

#v(1fr)

#align(right, block(width: auto, {{
  set par(first-line-indent: 0pt)
  text(size: 22pt, weight: 600, "{year}")
  linebreak()
  text(size: 14pt, "Volume {vol}, Number {num}")
  {f'linebreak(); v(0.4em); text(size: 11pt, style: "italic", fill: rgb("#4a4137"), {_typst_str(title)})' if title else ''}
}}))

#v(0.4in)

{f'''#pagebreak()

#align(center, text(size: 14pt, tracking: 0.18em, "CONTENTS"))
#v(1.5em)

#{{
  {toc_typst}
}}
''' if toc_typst else ''}
"""

    cover_typ = out_dir / "_cover.typ"
    cover_pdf = out_dir / "_cover.pdf"
    cover_typ.write_text(cover_typst, encoding="utf-8")
    typst_lib.compile(str(cover_typ), output=str(cover_pdf), root=str(CONTENT_DIR))
    return cover_pdf


def _short_journal_name(name: str) -> str:
    """Best-effort short form of a journal name for running headers.
    For 'Literacy in Composition Studies' -> 'LiCS'. Falls back to first
    letters of capitalized words."""
    if not name:
        return ""
    # If the name has a parenthetical acronym, prefer it.
    import re
    m = re.search(r"\(([A-Z][A-Za-z]{1,8})\)", name)
    if m:
        return m.group(1)
    # Otherwise use initials of words >2 chars (skip prepositions etc).
    skip = {"in", "of", "the", "and", "for", "on", "to", "a"}
    parts = [w for w in re.split(r"\s+", name) if w]
    initials = "".join(w[0] for w in parts if w.lower() not in skip)
    return initials or name


def _typst_str(s) -> str:
    """Encode a Python string as a Typst string literal."""
    if s is None:
        s = ""
    return '"' + str(s).replace("\\", "\\\\").replace('"', '\\"') + '"'


def _typst_inline_md(s: str) -> str:
    """Best-effort inline markdown -> Typst content for *italic* spans.
    Returns content that goes inside [...]."""
    if not s:
        return ""
    out = []
    i = 0
    while i < len(s):
        c = s[i]
        if c == "*" and i + 1 < len(s) and s[i + 1] != "*":
            end = s.find("*", i + 1)
            if end != -1:
                inner = s[i + 1:end].replace('"', '\\"')
                out.append(f'#text(style: "italic", "{inner}")')
                i = end + 1
                continue
        if c == '"':
            out.append('\\"')
        elif c == "\\":
            out.append("\\\\")
        else:
            out.append(c)
        i += 1
    return "".join(out)


def _typst_wordmark_block(rel_path: Optional[str], out_dir: Path, fallback: str) -> str:
    """Emit Typst code that renders the journal wordmark, either as an
    image (if the file exists) or as a large text fallback. The path is
    resolved via Typst's --root setting, so wordmark image paths are
    emitted as absolute-from-root references ("/journals/.../...")."""
    if rel_path:
        candidate = (CONTENT_DIR / rel_path).resolve()
        if candidate.exists():
            typst_path = "/" + rel_path.lstrip("/").replace("\\", "/")
            return (
                f'\n#v(0.7in)\n'
                f'#align(center, block(width: 100%, '
                f'image({_typst_str(typst_path)}, width: 75%, fit: "contain")))\n'
            )
    return (
        f'\n#v(1.5in)\n'
        f'#align(center, text(size: 96pt, weight: 700, font: ("Helvetica Neue", "Helvetica", "Arial"), {_typst_str(fallback)}))\n'
        f'#align(center, text(size: 18pt, font: ("Helvetica Neue", "Helvetica", "Arial"), {_typst_str("Literacy in Composition Studies" if fallback == "LiCS" else "")}))\n'
    )
